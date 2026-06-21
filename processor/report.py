"""报告生成模块

提供 Word 文档生成、合成图创建等功能。
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Inches

from processor.image_utils import resize_image_for_word, create_combined_image as _create_combined_image


def create_combined_image(hairstyle_path, user_path, result_paths, output_path):
    """创建合成图（委托给 image_utils）"""
    return _create_combined_image(hairstyle_path, user_path, result_paths, output_path)


def create_word_document(results, output_path="hairstyle_results.docx"):
    """使用所有结果创建 Word 文档"""
    doc = Document()
    doc.add_heading('发型换装结果', 0)
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'总共处理: {len(results)} 个组合')

    for i, result in enumerate(results):
        doc.add_heading(
            f'结果 {i+1}: {result["gender"]} - {result["user_filename"]} + {result["hairstyle_filename"]}',
            level=1
        )

        # 合成图
        if result.get('combined_image') and os.path.exists(result['combined_image']):
            doc.add_paragraph('拼接图片 (发型参考 + 用户照片 + 生成结果):')
            width, height = resize_image_for_word(result['combined_image'], max_width=6.0)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(result['combined_image'], width=Inches(width), height=Inches(height))
            doc.add_paragraph()

        # 单独图片表格
        doc.add_paragraph('单独图片:')
        result_images = result.get('result_images', [])
        num_cols = 2 + len(result_images)
        table = doc.add_table(rows=2, cols=num_cols)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '发型参考图'
        hdr_cells[1].text = '用户照片'
        for j in range(len(result_images)):
            hdr_cells[2 + j].text = f'生成结果{j+1}'

        img_cells = table.rows[1].cells

        def add_img_to_cell(cell, img_path):
            if os.path.exists(img_path):
                w, h = resize_image_for_word(img_path)
                p = cell.paragraphs[0]
                r = p.runs[0] if p.runs else p.add_run()
                r.add_picture(img_path, width=Inches(w), height=Inches(h))

        add_img_to_cell(img_cells[0], result['hairstyle_image'])
        add_img_to_cell(img_cells[1], result['user_image'])
        for j, result_image in enumerate(result_images):
            add_img_to_cell(img_cells[2 + j], result_image)

        doc.add_page_break()

    doc.save(output_path)
    print(f"Word 文档已保存: {output_path}")
